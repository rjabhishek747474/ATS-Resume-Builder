"""
Resume Download/Export API

Generates formatted resume output in PDF and DOCX formats.
"""

import io
from fastapi import APIRouter, HTTPException
from fastapi.responses import StreamingResponse
from pydantic import BaseModel
from typing import Optional

# Document generation
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer

from api.resume import resume_store
from api.jd import jd_store
from api.optimize import job_store

router = APIRouter()


class ResumeOutputRequest(BaseModel):
    job_id: str
    format: Optional[str] = "pdf"  # pdf or docx


@router.post("/resume/download")
async def download_resume(request: ResumeOutputRequest):
    """
    Generate downloadable resume in PDF or DOCX format.
    """
    if request.job_id not in job_store:
        raise HTTPException(404, "Job not found")
    
    job = job_store[request.job_id]
    
    if job.get("status") != "completed":
        raise HTTPException(400, f"Job not completed. Status: {job.get('status')}")
    
    result = job.get("result", {})
    optimized_resume = result.get("optimized_resume", {})
    ats_score = result.get("ats_score", 0)
    
    if request.format == "docx":
        # Generate DOCX
        doc_bytes = _generate_docx(optimized_resume, ats_score)
        return StreamingResponse(
            io.BytesIO(doc_bytes),
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": "attachment; filename=optimized_resume.docx"}
        )
    else:
        # Generate PDF (default)
        pdf_bytes = _generate_pdf(optimized_resume, ats_score)
        return StreamingResponse(
            io.BytesIO(pdf_bytes),
            media_type="application/pdf",
            headers={"Content-Disposition": "attachment; filename=optimized_resume.pdf"}
        )


@router.get("/resume/{resume_id}/formatted")
async def get_formatted_resume(resume_id: str, jd_id: Optional[str] = None):
    """
    Get formatted resume with optional JD keyword highlighting.
    """
    if resume_id not in resume_store:
        raise HTTPException(404, "Resume not found")
    
    resume_data = resume_store[resume_id]
    sections = resume_data.get("sections", {})
    
    # If JD provided, get keywords for highlighting
    keywords = []
    if jd_id and jd_id in jd_store:
        jd_data = jd_store[jd_id]
        keywords = jd_data.get("keywords", {}).get("primary", [])
        keywords += jd_data.get("hard_skills", [])
    
    formatted = _format_resume_with_keywords(sections, keywords)
    
    return {
        "resume_id": resume_id,
        "formatted_resume": formatted,
        "matched_keywords": keywords
    }


def _generate_pdf(sections: dict, ats_score: int) -> bytes:
    """Generate a professional PDF resume."""
    buffer = io.BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=letter, 
                           rightMargin=72, leftMargin=72,
                           topMargin=72, bottomMargin=72)
    
    styles = getSampleStyleSheet()
    
    # Custom styles
    title_style = ParagraphStyle(
        'CustomTitle',
        parent=styles['Heading1'],
        fontSize=16,
        spaceAfter=12,
        textColor='#1a365d'
    )
    
    section_style = ParagraphStyle(
        'SectionHeader',
        parent=styles['Heading2'],
        fontSize=12,
        spaceBefore=16,
        spaceAfter=8,
        textColor='#2c5282',
        borderColor='#2c5282',
        borderWidth=1,
        borderPadding=4
    )
    
    body_style = ParagraphStyle(
        'BodyText',
        parent=styles['Normal'],
        fontSize=10,
        spaceBefore=4,
        spaceAfter=4,
        leading=14
    )
    
    score_style = ParagraphStyle(
        'ScoreStyle',
        parent=styles['Normal'],
        fontSize=9,
        textColor='#718096',
        alignment=2  # Right align
    )
    
    story = []
    
    # Header with ATS Score
    story.append(Paragraph(f"ATS Compatibility Score: {ats_score}/100", score_style))
    story.append(Spacer(1, 12))
    
    # Section order
    section_order = ["summary", "experience", "skills", "education", "projects", "certifications"]
    section_titles = {
        "summary": "PROFESSIONAL SUMMARY",
        "experience": "PROFESSIONAL EXPERIENCE", 
        "skills": "TECHNICAL SKILLS",
        "education": "EDUCATION",
        "projects": "PROJECTS",
        "certifications": "CERTIFICATIONS"
    }
    
    for section_name in section_order:
        if section_name in sections and sections[section_name]:
            content = sections[section_name]
            
            # Section header
            title = section_titles.get(section_name, section_name.upper())
            story.append(Paragraph(title, section_style))
            
            # Section content - handle line breaks
            lines = content.split('\n')
            for line in lines:
                line = line.strip()
                if line:
                    # Escape special characters for PDF
                    line = line.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')
                    if line.startswith('-'):
                        line = '• ' + line[1:].strip()
                    story.append(Paragraph(line, body_style))
            
            story.append(Spacer(1, 8))
    
    doc.build(story)
    return buffer.getvalue()


def _generate_docx(sections: dict, ats_score: int) -> bytes:
    """Generate a professional DOCX resume."""
    doc = Document()
    
    # Set narrow margins
    for section in doc.sections:
        section.top_margin = Inches(0.75)
        section.bottom_margin = Inches(0.75)
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)
    
    # ATS Score header
    score_para = doc.add_paragraph()
    score_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    score_run = score_para.add_run(f"ATS Score: {ats_score}/100")
    score_run.font.size = Pt(9)
    score_run.font.italic = True
    
    # Section order
    section_order = ["summary", "experience", "skills", "education", "projects", "certifications"]
    section_titles = {
        "summary": "PROFESSIONAL SUMMARY",
        "experience": "PROFESSIONAL EXPERIENCE",
        "skills": "TECHNICAL SKILLS", 
        "education": "EDUCATION",
        "projects": "PROJECTS",
        "certifications": "CERTIFICATIONS"
    }
    
    for section_name in section_order:
        if section_name in sections and sections[section_name]:
            content = sections[section_name]
            
            # Section header
            title = section_titles.get(section_name, section_name.upper())
            heading = doc.add_heading(title, level=2)
            heading.runs[0].font.size = Pt(11)
            
            # Section content
            lines = content.split('\n')
            for line in lines:
                line = line.strip()
                if line:
                    if line.startswith('-'):
                        # Bullet point
                        para = doc.add_paragraph(line[1:].strip(), style='List Bullet')
                        para.runs[0].font.size = Pt(10)
                    else:
                        # Regular paragraph
                        para = doc.add_paragraph(line)
                        para.runs[0].font.size = Pt(10)
    
    # Save to bytes
    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


def _format_resume_with_keywords(sections: dict, keywords: list) -> str:
    """
    Format resume with keyword highlights.
    Keywords in the resume are marked with asterisks.
    """
    import re
    lines = []
    
    section_order = ["summary", "experience", "skills", "education"]
    
    for section_name in section_order:
        if section_name in sections and sections[section_name]:
            content = sections[section_name]
            
            # Highlight keywords (case-insensitive)
            for keyword in keywords:
                pattern = re.compile(re.escape(keyword), re.IGNORECASE)
                content = pattern.sub(f"**{keyword}**", content)
            
            lines.append(f"## {section_name.title()}")
            lines.append(content)
            lines.append("")
    
    return "\n".join(lines)
